import os
import logging
import json
from typing import List, Dict, Tuple, Optional, Any, Iterator
import shutil
import re
import time
import requests
import zipfile
import tempfile
import gdown

import torch
from sentence_transformers import SentenceTransformer
from pypdf import PdfReader
import docx as python_docx

from llama_index.core.llms import ChatMessage
from llama_index.llms.groq import Groq as LlamaIndexGroqClient

from langchain_groq import ChatGroq
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import ChatPromptTemplate
from langchain.schema import Document, BaseRetriever
from langchain.callbacks.manager import CallbackManagerForRetrieverRun
from langchain.schema.runnable import RunnablePassthrough, RunnableParallel
from langchain.schema.output_parser import StrOutputParser
from langchain.text_splitter import RecursiveCharacterTextSplitter
# MODIFIED: Import the new prompt
from system_prompts import RAG_SYSTEM_PROMPT, FALLBACK_SYSTEM_PROMPT, QA_FORMATTER_PROMPT

logger = logging.getLogger(__name__)
if not logger.handlers:
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )

# --- Configuration Constants ---
GROQ_API_KEY = os.getenv('BOT_API_KEY')
if not GROQ_API_KEY:
    logger.critical("CRITICAL: BOT_API_KEY environment variable not found. Services will fail.")

FALLBACK_LLM_MODEL_NAME = os.getenv("GROQ_FALLBACK_MODEL", "llama-3.1-70b-versatile")
# ADDED: New constant for the auxiliary model
AUXILIARY_LLM_MODEL_NAME = os.getenv("GROQ_AUXILIARY_MODEL", "llama3-8b-8192")
_MODULE_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
RAG_FAISS_INDEX_SUBDIR_NAME = "faiss_index"
RAG_STORAGE_PARENT_DIR = os.getenv("RAG_STORAGE_DIR", os.path.join(_MODULE_BASE_DIR, "faiss_storage"))
RAG_SOURCES_DIR = os.getenv("SOURCES_DIR", os.path.join(_MODULE_BASE_DIR, "sources"))
RAG_CHUNKED_SOURCES_FILENAME = "pre_chunked_sources.json"
os.makedirs(RAG_SOURCES_DIR, exist_ok=True)
os.makedirs(RAG_STORAGE_PARENT_DIR, exist_ok=True)
RAG_EMBEDDING_MODEL_NAME = os.getenv("RAG_EMBEDDING_MODEL", "all-MiniLM-L6-v2")
RAG_EMBEDDING_USE_GPU = os.getenv("RAG_EMBEDDING_GPU", "False").lower() == "true"
RAG_LLM_MODEL_NAME = os.getenv("RAG_LLM_MODEL", "llama-3.1-70b-versatile")
RAG_LLM_TEMPERATURE = float(os.getenv("RAG_TEMPERATURE", 0.0))
RAG_LOAD_INDEX_ON_STARTUP = os.getenv("RAG_LOAD_INDEX", "True").lower() == "true"
RAG_DEFAULT_RETRIEVER_K = int(os.getenv("RAG_RETRIEVER_K", 3))
GDRIVE_SOURCES_ENABLED = os.getenv("GDRIVE_SOURCES_ENABLED", "False").lower() == "true"
GDRIVE_FOLDER_ID_OR_URL = os.getenv("GDRIVE_FOLDER_URL")

# --- Text Extraction Helper Function ---
def extract_text_from_file(file_path: str, file_type: str) -> Optional[str]:
    logger.info(f"Extracting text from {file_type.upper()} file: {os.path.basename(file_path)}")
    try:
        if file_type == 'pdf':
            reader = PdfReader(file_path)
            return "".join(page.extract_text() + "\n" for page in reader.pages if page.extract_text())
        elif file_type == 'docx':
            doc = python_docx.Document(file_path)
            return "\n".join(para.text for para in doc.paragraphs if para.text)
        elif file_type == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        logger.warning(f"Unsupported file type for text extraction: {file_type}")
        return None
    except Exception as e:
        logger.error(f"Error extracting text from {os.path.basename(file_path)}: {e}", exc_info=True)
        return None

FAISS_RAG_SUPPORTED_EXTENSIONS = {'pdf': 'pdf', 'docx': 'docx', 'txt': 'txt'}

# --- FAISS RAG System ---
class FAISSRetrieverWithScore(BaseRetriever):
    vectorstore: FAISS
    k: int = RAG_DEFAULT_RETRIEVER_K

    def _get_relevant_documents(
        self, query: str, *, run_manager: CallbackManagerForRetrieverRun
    ) -> List[Document]:
        docs_and_scores = self.vectorstore.similarity_search_with_score(query, k=self.k)
        relevant_docs = []
        for doc, score in docs_and_scores:
            doc.metadata["retrieval_score"] = float(score)
            relevant_docs.append(doc)
        return relevant_docs

class KnowledgeRAG:
    def __init__(
        self,
        index_storage_dir: str,
        embedding_model_name: str,
        groq_model_name_for_rag: str,
        use_gpu_for_embeddings: bool,
        groq_api_key_for_rag: str,
        temperature: float,
    ):
        self.logger = logging.getLogger(__name__ + ".KnowledgeRAG")
        self.index_storage_dir = index_storage_dir
        self.embedding_model_name = embedding_model_name
        self.groq_model_name = groq_model_name_for_rag
        self.temperature = temperature

        device = "cuda" if use_gpu_for_embeddings and torch.cuda.is_available() else "cpu"
        self.logger.info(f"Initializing Hugging Face embedding model: {self.embedding_model_name} on device: {device}")
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.embedding_model_name,
                model_kwargs={"device": device},
                encode_kwargs={"normalize_embeddings": True}
            )
        except Exception as e:
            self.logger.critical(f"Failed to load embedding model: {e}", exc_info=True)
            raise

        self.logger.info(f"Initializing Langchain ChatGroq LLM for RAG: {self.groq_model_name}")
        if not groq_api_key_for_rag:
            raise ValueError("Groq API Key for RAG is missing.")
        try:
            self.llm = ChatGroq(
                temperature=self.temperature,
                groq_api_key=groq_api_key_for_rag,
                model_name=self.groq_model_name
            )
        except Exception as e:
            self.logger.critical(f"Failed to initialize Langchain ChatGroq LLM: {e}", exc_info=True)
            raise

        self.vector_store: Optional[FAISS] = None
        self.retriever: Optional[FAISSRetrieverWithScore] = None
        self.rag_chain = None
        self.processed_source_files: List[str] = []

    def build_index_from_source_files(self, source_folder_path: str, k: int = RAG_DEFAULT_RETRIEVER_K):
        all_docs_for_vectorstore: List[Document] = []
        processed_files_this_build: List[str] = []
        pre_chunked_json_path = os.path.join(self.index_storage_dir, RAG_CHUNKED_SOURCES_FILENAME)

        if os.path.exists(pre_chunked_json_path):
            self.logger.info(f"Loading documents from pre-chunked file: {pre_chunked_json_path}")
            try:
                with open(pre_chunked_json_path, 'r', encoding='utf-8') as f:
                    chunk_data_list = json.load(f)
                source_filenames = set()
                for chunk_data in chunk_data_list:
                    doc = Document(page_content=chunk_data.get("page_content", ""), metadata=chunk_data.get("metadata", {}))
                    all_docs_for_vectorstore.append(doc)
                    if 'source_document_name' in doc.metadata:
                        source_filenames.add(doc.metadata['source_document_name'])
                processed_files_this_build = sorted(list(source_filenames))
            except Exception as e:
                self.logger.error(f"Error processing pre-chunked JSON, falling back to raw files: {e}")
                all_docs_for_vectorstore.clear()

        if not all_docs_for_vectorstore:
            self.logger.info(f"Processing raw files from '{source_folder_path}' to build index.")
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=150)
            for filename in os.listdir(source_folder_path):
                file_path = os.path.join(source_folder_path, filename)
                file_ext = filename.split('.')[-1].lower()
                if os.path.isfile(file_path) and file_ext in FAISS_RAG_SUPPORTED_EXTENSIONS:
                    text_content = extract_text_from_file(file_path, file_ext)
                    if text_content:
                        chunks = text_splitter.split_text(text_content)
                        for i, chunk_text in enumerate(chunks):
                            metadata = {"source_document_name": filename, "chunk_index": i}
                            all_docs_for_vectorstore.append(Document(page_content=chunk_text, metadata=metadata))
                        processed_files_this_build.append(filename)

        if not all_docs_for_vectorstore:
            self.logger.warning(f"No processable PDF/DOCX/TXT documents found in '{source_folder_path}'. RAG index will only contain other sources if available.")


        self.processed_source_files = processed_files_this_build
        
        # This print statement is kept for console visibility on startup/rebuild
        print("\n--- Document Files Used for RAG Index ---")
        if self.processed_source_files:
            for filename in self.processed_source_files:
                print(f"- {filename}")
        else:
            print("No PDF/DOCX/TXT source files were processed for the RAG index.")
        print("---------------------------------------\n")

        if not all_docs_for_vectorstore:
            self.logger.warning("No documents to build FAISS index from. Skipping FAISS build.")
            return

        self.logger.info(f"Creating FAISS index from {len(all_docs_for_vectorstore)} document chunks...")
        self.vector_store = FAISS.from_documents(all_docs_for_vectorstore, self.embeddings)
        faiss_index_path = os.path.join(self.index_storage_dir, RAG_FAISS_INDEX_SUBDIR_NAME)
        self.vector_store.save_local(faiss_index_path)
        self.logger.info(f"FAISS index built and saved to '{faiss_index_path}'.")
        self.retriever = FAISSRetrieverWithScore(vectorstore=self.vector_store, k=k)
        self.setup_rag_chain()

    def load_index_from_disk(self, k: int = RAG_DEFAULT_RETRIEVER_K):
        faiss_index_path = os.path.join(self.index_storage_dir, RAG_FAISS_INDEX_SUBDIR_NAME)
        if not os.path.isdir(faiss_index_path):
            raise FileNotFoundError(f"FAISS index directory not found at '{faiss_index_path}'.")

        self.logger.info(f"Loading FAISS index from: {faiss_index_path}")
        self.vector_store = FAISS.load_local(
            folder_path=faiss_index_path,
            embeddings=self.embeddings,
            allow_dangerous_deserialization=True
        )
        self.retriever = FAISSRetrieverWithScore(vectorstore=self.vector_store, k=k)
        self.setup_rag_chain()

    def format_docs(self, docs: List[Document]) -> str:
        return "\n\n---\n\n".join([f"[Excerpt from {doc.metadata.get('source_document_name', 'N/A')}, Chunk {doc.metadata.get('chunk_index', 'N/A')}]\nContent:\n{doc.page_content}" for doc in docs])

    def setup_rag_chain(self):
        if not self.retriever or not self.llm:
            raise RuntimeError("Retriever and LLM must be initialized.")
        
        prompt = ChatPromptTemplate.from_template(RAG_SYSTEM_PROMPT)

        self.rag_chain = (
            RunnableParallel(
                context=(self.retriever | self.format_docs),
                question=RunnablePassthrough()
            )
            | prompt
            | self.llm
            | StrOutputParser()
        )
        self.logger.info("RAG LCEL chain set up successfully with dental assistant persona.")

    def invoke(self, query: str, top_k: Optional[int] = None) -> Dict[str, Any]:
        if not self.rag_chain:
            # MODIFIED: Changed severity
            self.logger.warning("RAG system not fully initialized. Cannot invoke.")
            return {"answer": "The provided bibliography does not contain specific information on this topic.", "source": "system_error", "cited_source_details": []}

        if not query or not query.strip():
            return {"answer": "Please provide a valid question.", "source": "system_error", "cited_source_details": []}

        k_to_use = top_k if top_k is not None and top_k > 0 else self.retriever.k
        self.logger.info(f"Processing RAG query with k={k_to_use}: '{query[:100]}...'")

        original_k = self.retriever.k
        if k_to_use != original_k:
            self.retriever.k = k_to_use
        
        try:
            retrieved_docs = self.retriever.get_relevant_documents(query)
            if not retrieved_docs:
                return {"answer": "The provided bibliography does not contain specific information on this topic.", "source": "no_docs_found", "cited_source_details": []}

            context_str = self.format_docs(retrieved_docs)
            
            # MODIFIED: Added full logging as per user request
            print(f"\n--- RAG INVOKE ---")
            print(f"QUESTION: {query}")
            print(f"CONTEXT:\n{context_str}")
            
            llm_answer = self.rag_chain.invoke(query, config={"context": context_str})

            print(f"LLM_ANSWER: {llm_answer}")
            print(f"--------------------\n")

            structured_sources = [{
                "source_document_name": doc.metadata.get('source_document_name', 'Unknown'),
                "chunk_index": doc.metadata.get('chunk_index', 'N/A'),
                "retrieval_score": doc.metadata.get("retrieval_score"),
            } for doc in retrieved_docs]

            if "the provided bibliography does not contain specific information" in llm_answer.lower():
                 final_answer = llm_answer
                 source_tag = "no_answer_in_bibliography"
            else:
                 final_answer = f"{llm_answer}\n\n*Source: Bibliography-Based*"
                 source_tag = "bibliography_based"

            return {
                "query": query,
                "answer": final_answer.strip(),
                "source": source_tag,
                "cited_source_details": structured_sources,
            }

        except Exception as e:
            self.logger.error(f"Error during RAG query processing: {e}", exc_info=True)
            return {"answer": "An error occurred while processing your request.", "source": "system_error", "cited_source_details": []}
        finally:
            if k_to_use != original_k:
                self.retriever.k = original_k
    
    def stream(self, query: str, top_k: Optional[int] = None) -> Iterator[str]:
        if not self.rag_chain:
            self.logger.error("RAG system not fully initialized for streaming.")
            yield "Error: RAG system is not ready."
            return

        k_to_use = top_k if top_k is not None and top_k > 0 else self.retriever.k
        self.logger.info(f"Processing RAG stream with k={k_to_use}: '{query[:100]}...'")

        original_k = self.retriever.k
        if k_to_use != original_k:
            self.retriever.k = k_to_use
        
        try:
            # Check for docs first to avoid streaming "no info" message
            retrieved_docs = self.retriever.get_relevant_documents(query)
            if not retrieved_docs:
                yield "The provided bibliography does not contain specific information on this topic."
                return
            
            # MODIFIED: Added full logging for streaming as per user request
            context_str = self.format_docs(retrieved_docs)
            print(f"\n--- RAG STREAM ---")
            print(f"QUESTION: {query}")
            print(f"CONTEXT:\n{context_str}")
            print(f"STREAMING LLM_ANSWER...")
            print(f"--------------------\n")
                
            yield from self.rag_chain.stream(query, config={"context": context_str})
        except Exception as e:
            self.logger.error(f"Error during RAG stream processing: {e}", exc_info=True)
            yield "An error occurred while processing your request."
        finally:
            if k_to_use != original_k:
                self.retriever.k = original_k


# --- Groq Fallback Bot ---
class GroqBot:
    def __init__(self):
        self.logger = logging.getLogger(__name__ + ".GroqBot")
        if not GROQ_API_KEY:
            self.client = None
            self.logger.critical("GroqBot not initialized: BOT_API_KEY is missing.")
            return
        try:
            self.client = LlamaIndexGroqClient(model=FALLBACK_LLM_MODEL_NAME, api_key=GROQ_API_KEY)
            self.system_prompt = FALLBACK_SYSTEM_PROMPT
        except Exception as e:
            self.logger.error(f"Failed to initialize LlamaIndexGroqClient for Fallback Bot: {e}", exc_info=True)
            self.client = None

    def stream_response(self, context: dict) -> Iterator[str]:
        if not self.client:
            yield "The system is currently unable to process this request."
            return

        current_query = context.get('current_query', '')
        chat_history = context.get('chat_history', [])
        qa_info = context.get('qa_related_info', '')

        messages = [ChatMessage(role="system", content=self.system_prompt)]
        if chat_history:
            messages.extend([ChatMessage(**msg) for msg in chat_history])
        if qa_info:
            messages.append(ChatMessage(role="system", content=f"**Potentially Relevant Q&A Information from other sources:**\n{qa_info}"))
        messages.append(ChatMessage(role="user", content=f"**Current User Query:**\n{current_query}"))
        
        # MODIFIED: Added full logging as per user request
        # The conversion to dict is necessary because ChatMessage is not directly JSON serializable
        messages_for_print = [msg.dict() for msg in messages]
        print(f"\n--- FALLBACK STREAM ---")
        print(f"MESSAGES SENT TO LLM:\n{json.dumps(messages_for_print, indent=2)}")
        print(f"STREAMING LLM_ANSWER...")
        print(f"-----------------------\n")

        try:
            response_stream = self.client.stream_chat(messages)
            for r_chunk in response_stream:
                yield r_chunk.delta
        except Exception as e:
            self.logger.error(f"Groq API error in get_response (Fallback): {e}", exc_info=True)
            yield "I am currently unable to process this request due to a technical issue."

# ADDED: New function for formatting QA answers
def get_answer_from_context(question: str, context: str, system_prompt: str) -> str:
    """
    Calls the LLM with a specific question and context from a QA source (CSV/XLSX).
    """
    logger.info(f"Formatting answer for question '{question[:50]}...' using QA context.")
    try:
        # Use the auxiliary model for this task for speed and cost-efficiency
        formatter_llm = ChatGroq(
            temperature=0.1,
            groq_api_key=GROQ_API_KEY,
            model_name=AUXILIARY_LLM_MODEL_NAME
        )
        
        prompt_template = ChatPromptTemplate.from_template(system_prompt)
        
        chain = prompt_template | formatter_llm | StrOutputParser()
        
        # MODIFIED: Added full logging as per user request
        print(f"\n--- QA FORMATTER ---")
        print(f"QUESTION: {question}")
        print(f"CONTEXT:\n{context}")

        response = chain.invoke({
            "context": context,
            "question": question
        })

        print(f"LLM_ANSWER: {response}")
        print(f"--------------------\n")
        
        return response.strip()
        
    except Exception as e:
        logger.error(f"Error in get_answer_from_context: {e}", exc_info=True)
        return "Sorry, I was unable to formulate an answer based on the available information."


# --- Initialization and Interface Functions ---
def get_id_from_gdrive_input(url_or_id: str) -> Optional[str]:
    if not url_or_id: return None
    patterns = [r"/folders/([a-zA-Z0-9_-]+)", r"/d/([a-zA-Z0-9_-]+)", r"id=([a-zA-Z0-9_-]+)"]
    for pattern in patterns:
        match = re.search(pattern, url_or_id)
        if match: return match.group(1)
    if "/" not in url_or_id and "=" not in url_or_id and len(url_or_id) > 15:
        return url_or_id
    return None

def download_and_unzip_gdrive_folder(folder_id_or_url: str, target_dir: str) -> bool:
    folder_id = get_id_from_gdrive_input(folder_id_or_url)
    if not folder_id:
        logger.error(f"Invalid Google Drive Folder ID or URL: {folder_id_or_url}")
        return False
    
    with tempfile.TemporaryDirectory() as temp_dir:
        try:
            logger.info(f"Attempting to download GDrive folder ID: {folder_id}")
            download_path = gdown.download_folder(id=folder_id, output=temp_dir, quiet=False, use_cookies=False)
            if not download_path or not os.listdir(temp_dir):
                logger.error("gdown failed to download or extract the folder.")
                return False

            source_content_root = temp_dir
            items_in_temp = os.listdir(temp_dir)
            if len(items_in_temp) == 1 and os.path.isdir(os.path.join(temp_dir, items_in_temp[0])):
                source_content_root = os.path.join(temp_dir, items_in_temp[0])
            
            logger.info(f"Moving contents from {source_content_root} to {target_dir}")
            if os.path.exists(target_dir):
                shutil.rmtree(target_dir)
            shutil.copytree(source_content_root, target_dir)
            logger.info(f"Successfully moved GDrive contents to {target_dir}")
            return True
        except Exception as e:
            # MODIFIED: Corrected self.logger to logger
            logger.error(f"Error during GDrive download/processing: {e}", exc_info=True)
            return False

def initialize_and_get_rag_system(force_rebuild: bool = False) -> Optional[KnowledgeRAG]:
    if not GROQ_API_KEY:
        logger.error("RAG system cannot be initialized without BOT_API_KEY.")
        return None

    if GDRIVE_SOURCES_ENABLED and GDRIVE_FOLDER_ID_OR_URL:
        logger.info("Google Drive sources enabled. Downloading...")
        if os.path.isdir(RAG_SOURCES_DIR):
            logger.info(f"Clearing existing RAG sources directory: {RAG_SOURCES_DIR}")
            shutil.rmtree(RAG_SOURCES_DIR)
            os.makedirs(RAG_SOURCES_DIR)
        
        download_successful = download_and_unzip_gdrive_folder(GDRIVE_FOLDER_ID_OR_URL, RAG_SOURCES_DIR)
        if not download_successful:
            logger.error("Failed to download sources from Google Drive. Using local files if available.")

    faiss_index_path = os.path.join(RAG_STORAGE_PARENT_DIR, RAG_FAISS_INDEX_SUBDIR_NAME)
    if force_rebuild and os.path.exists(RAG_STORAGE_PARENT_DIR):
        logger.info(f"Force Rebuild: Deleting existing index storage directory at '{RAG_STORAGE_PARENT_DIR}'")
        shutil.rmtree(RAG_STORAGE_PARENT_DIR)
        os.makedirs(RAG_STORAGE_PARENT_DIR)

    try:
        rag_instance = KnowledgeRAG(
            index_storage_dir=RAG_STORAGE_PARENT_DIR,
            embedding_model_name=RAG_EMBEDDING_MODEL_NAME,
            groq_model_name_for_rag=RAG_LLM_MODEL_NAME,
            use_gpu_for_embeddings=RAG_EMBEDDING_USE_GPU,
            groq_api_key_for_rag=GROQ_API_KEY,
            temperature=RAG_LLM_TEMPERATURE,
        )
        
        should_build = True
        if RAG_LOAD_INDEX_ON_STARTUP and not force_rebuild:
            try:
                rag_instance.load_index_from_disk(k=RAG_DEFAULT_RETRIEVER_K)
                logger.info("RAG index loaded successfully from disk.")
                should_build = False
            except FileNotFoundError:
                logger.warning("Index not found on disk. Will attempt to build.")
            except Exception as e:
                logger.error(f"Error loading index: {e}. Will attempt to rebuild.", exc_info=True)

        if should_build:
            logger.info("Building new RAG index from source data...")
            rag_instance.build_index_from_source_files(
                source_folder_path=RAG_SOURCES_DIR,
                k=RAG_DEFAULT_RETRIEVER_K
            )
        
        return rag_instance

    except Exception as e:
        logger.critical(f"FATAL: Failed to initialize RAG system: {e}", exc_info=True)
        return None

groq_bot_instance = GroqBot()

# ADDED: New function to handle auxiliary model calls (This function is no longer used, replaced by get_answer_from_context)
def get_auxiliary_chat_response(messages: List[Dict]) -> str:
    """
    Handles requests for auxiliary tasks like generating titles or follow-up questions.
    Uses a separate, smaller model for efficiency.
    """
    logger.info(f"Routing auxiliary request to model: {AUXILIARY_LLM_MODEL_NAME}")
    try:
        # Initialize a dedicated client for this call to use the specific auxiliary model
        aux_client = ChatGroq(
            temperature=0.2, # A bit more creative than RAG, but still grounded
            groq_api_key=GROQ_API_KEY,
            model_name=AUXILIARY_LLM_MODEL_NAME
        )
        response = aux_client.invoke(messages)
        return response.content
    except Exception as e:
        logger.error(f"Error with auxiliary model call: {e}", exc_info=True)
        return "Could not generate suggestions."